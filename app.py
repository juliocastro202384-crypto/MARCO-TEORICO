# app.py -- Constructor de Marco Teorico v7.5.0 (P-7 + Fix1: keyword filter + Fix2: noise filter + Fix3: DOI GET verification)
# Sidebar reestructurado: 6 secciones + boton Recuperar separado

import io
import re
import markdown

def md_to_html(text: str) -> str:
    """Convierte markdown a HTML usando la libreria markdown con soporte de tablas."""
    text = re.sub(r'\n{3,}', '\n\n', text)
    return markdown.markdown(
        text,
        extensions=['tables', 'fenced_code'],
    )
import math
import os
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
import streamlit as st
import anthropic
from docx import Document
from rapidfuzz import fuzz

ANIO_ACTUAL    = datetime.now().year
ANIO_INICIO    = ANIO_ACTUAL - 5
RANGO          = f"{ANIO_INICIO}-{ANIO_ACTUAL}"

OPENALEX         = "https://api.openalex.org"
CROSSREF         = "https://api.crossref.org"
S2               = "https://api.semanticscholar.org/graph/v1"
CROSSREF_WORKERS = 8
S2_FIELDS = ("title,year,venue,abstract,url,authors,citationCount,externalIds,publicationTypes")

st.set_page_config(
    page_title="Constructor de Marco Teorico - Claude",
    page_icon=":classical_building:",
    layout="wide",
)

CSS = """
<style>
  html, body, [class*="css"] { font-family: Georgia, serif; }
  [data-testid="stSidebar"] {
    background: linear-gradient(180deg, #f8f6f0 0%, #eef2f7 100%);
    border-right: 2px solid #2563eb;
  }
  [data-testid="stSidebar"] label { color: #374151; font-weight: 600; font-size: 0.85rem; }
  .main-header {
    background: linear-gradient(135deg, #1e3a5f 0%, #2563eb 50%, #3b82f6 100%);
    color: white; padding: 2rem 2.5rem; border-radius: 12px;
    margin-bottom: 1.5rem; box-shadow: 0 4px 20px rgba(37,99,235,0.3);
  }
  .main-header h1 { margin: 0; font-size: 1.8rem; font-weight: 700; }
  .main-header p { margin: 0.5rem 0 0; opacity: 0.9; font-size: 0.95rem; }
  .badge-row { display: flex; gap: 8px; margin-top: 0.8rem; flex-wrap: wrap; }
  .badge { background: rgba(255,255,255,0.2); border: 1px solid rgba(255,255,255,0.4); color: white; padding: 3px 10px; border-radius: 20px; font-size: 0.75rem; font-weight: 600; }
  .info-box { background: #eff6ff; border-left: 4px solid #2563eb; padding: 1rem 1.2rem; border-radius: 0 8px 8px 0; margin: 1rem 0; }
  .warning-box { background: #fffbeb; border-left: 4px solid #f59e0b; padding: 1rem 1.2rem; border-radius: 0 8px 8px 0; margin: 1rem 0; }
  .result-container { background: #fafafa; border: 1px solid #e5e7eb; border-radius: 10px; padding: 1.5rem 2rem; margin-top: 1rem; font-size: 0.95rem; line-height: 1.7; }
    .result-container h2 { font-size: 1.15rem; font-weight: 700; color: #1e3a5f; margin: 1.2rem 0 0.3rem; border-bottom: 1px solid #dbeafe; padding-bottom: 0.2rem; }
    .result-container h3 { font-size: 1.0rem; font-weight: 700; color: #2563eb; margin: 0.9rem 0 0.2rem; }
    .result-container p { margin: 0.35rem 0; text-align: justify; }
    .result-container strong { color: #111827; }
    .result-container table { border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: 0.9rem; }
    .result-container th { background-color: #2c3e50; color: white; padding: 8px 12px; text-align: left; }
    .result-container td { padding: 7px 12px; border-bottom: 1px solid #dee2e6; }
    .result-container tr:nth-child(even) td { background-color: #f8f9fa; }
  .stButton > button { background: linear-gradient(135deg, #1e3a5f, #2563eb); color: white; border: none; border-radius: 8px; padding: 0.6rem 1.2rem; font-weight: 600; width: 100%; }
  footer { color: #9ca3af; font-size: 0.8rem; text-align: center; margin-top: 2rem; padding-top: 1rem; border-top: 1px solid #e5e7eb; }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

_GENERIC_TERMS = {"the","and","for","that","this","with","from","como","para","que","los","las","del","una","por"}

def norm(text: str) -> str:
    return re.sub(r"[^a-z0-9]", "", text.lower().strip())

def _relevance_score(text: str, term: str) -> float:
    tn = norm(term)
    if not tn or tn in _GENERIC_TERMS or len(tn) < 5:
        return 0.0
    words = text.split()
    if not words:
        return 0.0
    hits = sum(1 for w in words if tn in w)
    if hits == 0:
        return 0.0
    tf = hits / len(words)
    length_penalty = min(len(words) / 50.0, 1.0)
    return min(tf * length_penalty * 10, 1.0)

def is_header_like(p: str) -> bool:
    parts = p.split("\n")
    if len(parts) > 2:
        return False
    first = parts[0].strip()
    if re.match(r"^#{1,4}\s", first):
        return True
    if first == first.upper() and first != first.lower() and len(first) <= 80:
        return True
    if re.match(r"^(\d{1,2}|[IVX]{1,4})[.)]\s+[A-Z]", first) and len(first) <= 80:
        return True
    if first.endswith(":") and len(first) <= 60:
        return True
    return False

def is_valid_doi(doi: str) -> bool:
    if not doi:
        return False
    d = doi.strip().lower()
    return d.startswith("10.") and "/" in d

# ===== FIX 1: Filtro de keywords irrelevantes por disciplina =====
_IRRELEVANT_KEYWORDS = [
    "matemáticas", "matematicas", "física", "fisica",
    "conducta", "tea", "instagram",
]

def is_relevant_record(record):
    """Retorna False si el titulo o abstract contiene keywords irrelevantes."""
    text = (
        (record.get("title") or "") + " " + (record.get("abstract") or "")
    ).lower()
    for kw in _IRRELEVANT_KEYWORDS:
        if kw in text:
            return False
    return True

# ===== FIX 2: Excluir titulos <4 palabras o tipo book-section =====
def is_noise_record(record):
    """Retorna True si el registro es ruido bibliografico."""
    title = (record.get("title") or "").strip()
    word_count = len(title.split())
    if word_count < 4:
        return True
    pub_types = record.get("publication_types") or record.get("publicationTypes") or []
    if isinstance(pub_types, list):
        for pt in pub_types:
            if isinstance(pt, str) and "book-section" in pt.lower():
                return True
    elif isinstance(pub_types, str):
        if "book-section" in pub_types.lower():
            return True
    return False


# ===== P-7: Traduccion, Queries especificas, Clasificador, Fallback =====

# Diccionario de traduccion ES→EN para terminos comunes en ciencias sociales
_ES_EN = {
    "competencias digitales": "digital competencies",
    "competencia digital": "digital competency",
    "tecnologia educativa": "educational technology",
    "rendimiento academico": "academic performance",
    "aprendizaje": "learning",
    "ensenanza": "teaching",
    "docentes": "teachers",
    "estudiantes": "students",
    "motivacion": "motivation",
    "liderazgo": "leadership",
    "gestion": "management",
    "clima organizacional": "organizational climate",
    "satisfaccion laboral": "job satisfaction",
    "desempeno": "performance",
    "innovacion": "innovation",
    "inclusion": "inclusion",
    "educacion": "education",
    "salud mental": "mental health",
    "bienestar": "wellbeing",
    "ansiedad": "anxiety",
    "resiliencia": "resilience",
    "autoeficacia": "self-efficacy",
    "inteligencia emocional": "emotional intelligence",
    "habilidades": "skills",
    "alfabetizacion": "literacy",
    "evaluacion": "assessment",
    "curriculo": "curriculum",
    "pedagogia": "pedagogy",
    "didactica": "didactics",
    "constructivismo": "constructivism",
    "aprendizaje significativo": "meaningful learning",
    "trabajo colaborativo": "collaborative work",
    "pensamiento critico": "critical thinking",
    "resolucion de problemas": "problem solving",
    "toma de decisiones": "decision making",
    "comunicacion": "communication",
    "entornos virtuales": "virtual environments",
    "educacion a distancia": "distance education",
    "e-learning": "e-learning",
    "blended learning": "blended learning",
    "gamificacion": "gamification",
    "inteligencia artificial": "artificial intelligence",
    "big data": "big data",
    "redes sociales": "social networks",
    "emprendimiento": "entrepreneurship",
    "sostenibilidad": "sustainability",
    "equidad": "equity",
    "genero": "gender",
    "interculturalidad": "interculturality",
    "violencia escolar": "school violence",
    "bullying": "bullying",
    "desercion escolar": "school dropout",
    "nivel socioeconomico": "socioeconomic status",
    "familia": "family",
    "padres": "parents",
    "comunidad": "community",
}

def translate_query(text: str) -> str:
    """Traduce terminos clave ES→EN para mejorar cobertura en APIs."""
    if not text:
        return text
    t = text.lower().strip()
    # Ordenar por longitud descendente para reemplazar frases antes que palabras sueltas
    for es, en in sorted(_ES_EN.items(), key=lambda x: -len(x[0])):
        t = t.replace(es, en)
    return t

def build_queries(vi: str, vd: str, titulo: str, area: str) -> List[str]:
    """Genera queries especificas a partir de VI/VD: ES + EN + combinadas."""
    queries = []
    vi_c = vi.strip() if vi else ""
    vd_c = vd.strip() if vd else ""
    titulo_c = titulo.strip() if titulo else ""

    # Query 1: VI + VD en español
    if vi_c and vd_c:
        queries.append(f"{vi_c} {vd_c}")
    elif vi_c:
        queries.append(vi_c)
    elif vd_c:
        queries.append(vd_c)

    # Query 2: traduccion al ingles
    vi_en = translate_query(vi_c)
    vd_en = translate_query(vd_c)
    if vi_en and vd_en and (vi_en != vi_c or vd_en != vd_c):
        queries.append(f"{vi_en} {vd_en}")
    elif vi_en and vi_en != vi_c:
        queries.append(vi_en)

    # Query 3: titulo del estudio
    if titulo_c and titulo_c not in queries:
        queries.append(titulo_c)

    # Query 4: VI sola en ingles
    if vi_en and vi_en not in queries:
        queries.append(vi_en)

    # Query 5: VD sola en ingles
    if vd_en and vd_en not in queries and vd_en != vi_en:
        queries.append(vd_en)

    # Eliminar duplicados y vacios, limitar a 4
    seen = set()
    final = []
    for q in queries:
        q = q.strip()
        if q and q not in seen:
            seen.add(q)
            final.append(q)
    return final[:4]

def classify_source(record: Dict, year_from: int) -> str:
    """Clasifica una fuente como conceptual, teorica o empirica."""
    title = (record.get("title") or "").lower()
    abstract = (record.get("abstract") or "").lower()
    year = record.get("year") or 0
    combined = title + " " + abstract

    # Indicadores empíricos
    empirical_kw = ["study", "survey", "sample", "participants", "n=", "results show",
                    "findings", "regression", "correlation", "experiment", "data",
                    "questionnaire", "analysis", "estudio", "muestra", "participantes",
                    "resultados", "hallazgos", "regresion", "correlacion", "datos",
                    "encuesta", "cuestionario", "analisis empirico"]
    # Indicadores teóricos
    theory_kw = ["theory", "theoretical", "framework", "model", "conceptual framework",
                 "teoria", "teorico", "marco conceptual", "modelo teorico", "paradigma",
                 "epistemolog", "ontolog", "review", "meta-analysis", "systematic review",
                 "revision sistematica", "metaanalisis"]

    emp_score = sum(1 for kw in empirical_kw if kw in combined)
    the_score = sum(1 for kw in theory_kw if kw in combined)

    if emp_score >= 2 and year >= year_from:
        return "empirica"
    elif the_score >= 2:
        return "teorica"
    elif emp_score >= 1 and year >= year_from:
        return "empirica"
    else:
        return "conceptual"

@st.cache_data(ttl=3600, show_spinner=False)
def crossref_search(query: str, year_from: int, max_results: int = 8) -> List[Dict]:
    """Busca en Crossref API por query y filtra por anio."""
    try:
        params = {
            "query": query,
            "rows": max_results,
            "filter": f"from-pub-date:{year_from}",
            "select": "DOI,title,author,published,container-title,abstract",
        }
        r = requests.get(
            f"{CROSSREF}/works",
            params=params,
            timeout=10,
            headers={"User-Agent": "MarcoTeoricoApp/7.0"},
        )
        if r.status_code != 200:
            return []
        results = []
        for item in r.json().get("message", {}).get("items", []):
            doi = (item.get("DOI") or "").strip()
            titles = item.get("title") or []
            title = titles[0] if titles else ""
            authors_raw = item.get("author") or []
            authors = [
                f"{a.get('family', '')} {a.get('given', '')[:1]}".strip()
                for a in authors_raw[:3]
            ]
            pub = item.get("published") or item.get("published-print") or {}
            dp = pub.get("date-parts") or [[None]]
            year = dp[0][0] if dp and dp[0] else None
            venue_list = item.get("container-title") or []
            venue = venue_list[0] if venue_list else ""
            abstract = (item.get("abstract") or "")[:400]
            abstract = re.sub(r"<[^>]+>", "", abstract)  # strip HTML
            if not title:
                continue
            results.append({
                "id": f"CR_{len(results)+1}",
                "source": "Crossref",
                "title": title,
                "year": year,
                "authors": authors,
                "venue": venue,
                "doi": doi,
                "abstract": abstract,
                "open_access": False,
                "verified_by": ["Crossref"],
                "quality_flags": {},
            })
        return results
    except Exception:
        return []

@st.cache_data(ttl=3600, show_spinner=False)
def crossref_ok(doi: str) -> bool:
    try:
        r = requests.get(f"{CROSSREF}/works/{doi.strip()}", timeout=8, headers={"User-Agent": "MarcoTeoricoApp/7.0"})
        return r.status_code == 200
    except Exception:
        return False

# ===== FIX 3: Verificacion de DOI via GET directo =====
@st.cache_data(ttl=3600, show_spinner=False)
def doi_verified(doi: str) -> bool:
    """Hace GET al DOI resolver (doi.org). Marca VERIFICADA solo si responde HTTP 200."""
    if not doi or not is_valid_doi(doi):
        return False
    doi_clean = doi.strip().lstrip("https://doi.org/").lstrip("http://doi.org/")
    try:
        r = requests.get(
            f"https://doi.org/{doi_clean}",
            timeout=10,
            allow_redirects=True,
            headers={"User-Agent": "MarcoTeoricoApp/7.0", "Accept": "application/json"},
        )
        return r.status_code == 200
    except Exception:
        return False

@st.cache_data(ttl=3600, show_spinner=False)
def openalex_search(query: str, variables: str, year_from: int, max_results: int = 10) -> List[Dict]:
    """OpenAlex: busca con query principal + traduccion EN. Filtra por anio."""
    try:
        query_en = translate_query(query)
        search_term = query_en if query_en != query else query
        params = {
            "search": search_term,
            "filter": f"from_publication_date:{year_from}-01-01,is_paratext:false",
            "per-page": max_results,
            "select": "id,doi,title,publication_year,authorships,primary_location,open_access,abstract_inverted_index,primary_topic",
        }
        r = requests.get(
            f"{OPENALEX}/works",
            params=params,
            timeout=10,
            headers={"User-Agent": "MarcoTeoricoApp/7.0"},
        )
        if r.status_code != 200:
            return []
        results = []
        for w in r.json().get("results", []):
            doi = (w.get("doi") or "").replace("https://doi.org/", "")
            authors = [
                a["author"]["display_name"]
                for a in w.get("authorships", [])[:3]
                if a.get("author")
            ]
            loc = w.get("primary_location") or {}
            venue = (loc.get("source") or {}).get("display_name", "")
            inv = w.get("abstract_inverted_index") or {}
            abstract = ""
            if inv:
                pos = sorted((p, wd) for wd, pl in inv.items() for p in pl)
                abstract = " ".join(wd for _, wd in pos[:80])
            results.append({
                "id": f"OA_{len(results)+1}",
                "source": "OpenAlex",
                "title": w.get("title", ""),
                "year": w.get("publication_year"),
                "authors": authors,
                "venue": venue,
                "doi": doi,
                "abstract": abstract,
                "open_access": (w.get("open_access") or {}).get("is_oa", False),
                "verified_by": [],
                "quality_flags": {},
            })
        return results
    except Exception:
        return []

@st.cache_data(ttl=3600, show_spinner=False)
def s2_search(query: str, variables: str, max_results: int = 8) -> List[Dict]:
    """Semantic Scholar: busca con traduccion EN automatica."""
    try:
        query_en = translate_query(query)
        search_term = query_en if query_en != query else query
        r = requests.get(
            f"{S2}/paper/search",
            params={"query": search_term, "limit": max_results, "fields": S2_FIELDS},
            timeout=10,
        )
        if r.status_code != 200:
            return []
        results = []
        for p in r.json().get("data", []):
            doi = (p.get("externalIds") or {}).get("DOI", "")
            authors = [a.get("name", "") for a in p.get("authors", [])[:3]]
            results.append({
                "id": f"S2_{len(results)+1}",
                "source": "SemanticScholar",
                "title": p.get("title", ""),
                "year": p.get("year"),
                "authors": authors,
                "venue": p.get("venue", ""),
                "doi": doi,
                "abstract": (p.get("abstract") or "")[:500],
                "open_access": False,
                "verified_by": [],
                "quality_flags": {},
            })
        return results
    except Exception:
        return []

def verify_records_concurrent(records: List[Dict], progress_bar=None) -> List[Dict]:
    doi_seen: set = set()
    unique_records: List[Dict] = []
    no_doi: List[Dict] = []
    for r in records:
        doi = (r.get("doi") or "").lower().strip()
        if doi and doi not in doi_seen:
            doi_seen.add(doi)
            unique_records.append(r)
        elif not doi:
            no_doi.append(r)
    all_to_check = unique_records + no_doi
    total = len(all_to_check)
    done = 0
    verified = []
    def _check(rec: Dict) -> Tuple[bool, Dict]:
        doi = rec.get("doi")
        # FIX 3: GET al DOI - VERIFICADA solo si responde 200
        doi_ok = bool(doi and is_valid_doi(doi) and doi_verified(doi))
        crossref_ok_flag = bool(doi and is_valid_doi(doi) and crossref_ok(doi))
        ok = doi_ok or crossref_ok_flag
        rec["verified_by"] = ["DOI-Verificado"] if doi_ok else (["Crossref"] if crossref_ok_flag else [])
        rec["quality_flags"]["has_doi"] = bool(doi)
        rec["quality_flags"]["has_venue"] = bool(rec.get("venue"))
        rec["quality_flags"]["has_year"] = bool(rec.get("year"))
        rec["quality_flags"]["has_authors"] = bool(rec.get("authors"))
        rec["quality_flags"]["has_abstract"] = bool(rec.get("abstract"))
        return ok, rec
    with ThreadPoolExecutor(max_workers=CROSSREF_WORKERS) as pool:
        futures = {pool.submit(_check, rec): rec for rec in all_to_check}
        for future in as_completed(futures):
            done += 1
            if progress_bar:
                progress_bar.progress(min(done / max(total, 1), 1.0))
            try:
                ok, rec = future.result()
                if ok:
                    verified.append(rec)
            except Exception:
                pass
    return verified

def dedup_records(records: List[Dict]) -> List[Dict]:
    seen_doi: set = set()
    seen_title: set = set()
    out = []
    for r in records:
        doi = (r.get("doi") or "").lower().strip()
        tn = norm(r.get("title") or "")
        if doi and doi in seen_doi:
            continue
        if any(fuzz.ratio(tn, t) > 88 for t in seen_title):
            continue
        if doi:
            seen_doi.add(doi)
        if tn:
            seen_title.add(tn)
        out.append(r)
    return out

def generar_docx(texto: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Marco Teorico", 0)
    for linea in texto.split("\n"):
        s = linea.strip()
        if not s:
            continue
        if re.match(r"^#{1,2}\s", s):
            doc.add_heading(re.sub(r"^#+\s+", "", s), level=1 if s.startswith("# ") else 2)
        elif re.match(r"^(I{1,3}|IV|V?I{0,3}|IX|X)\. \S", s, re.I):
            doc.add_heading(s, level=1)
        elif re.match(r"^\d{1,2}\. [A-Z]", s):
            doc.add_heading(s, level=2)
        else:
            doc.add_paragraph(s)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generar_marco_completo(variables, fuentes, system_prompt, client, modelo, result_area):
    """Llamada 1/2: genera S0-S6 con streaming. Retorna el texto completo."""

    # ══ LLAMADA 1: S0 al S6 ══
    prompt_1 = f"""
{system_prompt}

VARIABLES: {variables}

FUENTES VERIFICADAS: {fuentes}

INSTRUCCION: Genera UNICAMENTE las secciones S0, S1, S2, S3, S4, S5 y S6.
Termina exactamente al cerrar S6. No escribas S7 ni ninguna seccion posterior.
"""

    parte_1 = ""
    with client.messages.stream(
        model=modelo,
        max_tokens=6000,
        messages=[{"role": "user", "content": prompt_1}],
    ) as stream:
        for chunk in stream.text_stream:
            parte_1 += chunk
            result_area.markdown(
                f"<div class='result-container'>{md_to_html(parte_1)}"
                f"<p style='color:#6b7280;font-size:0.8rem;margin-top:0.8rem'>Generando S0-S6...</p></div>",
                unsafe_allow_html=True,
            )
    return parte_1


SYSTEM_PROMPT = (
    "AGENTE: CONSTRUCTOR DE MARCO TEORICO v7.0 (P-7)\n"
    "Eres un AGENTE ACADEMICO DE ALTO RIGOR DOCTORAL.\n\n"
    "0. DECISION DE MODO\n"
    "- Gate global cumplido: Si/No\n"
    "- Modo activado: MODO A o MODO B\n\n"
    "I. GATE GLOBAL\n"
    "- FALTA suficiencia en 1 variable -> MODO A\n"
    "- TODAS cumplen -> MODO B\n\n"
    "II. SUFICIENCIA MINIMA\n"
    "VERIFICADA: Autor+Anio+Titulo+Revista. Min: 2 conceptual, 1 teorico, 2 empiricos.\n"
    "Scholar sin DOI -> NO verificada.\n\n"
    "III. MODO A (5 secciones)\n"
    "S1: DECISION DE MODO\nS2: VACIOS\nS3: SEMILLAS\nS4: CONSULTAS\nS5: PLANTILLA\n"
    "ULTIMA LINEA MODO A: REINYECTAR FUENTES_RECUPERADAS PARA ACTIVAR MODO B.\n"
    "PROHIBIDO MODO A: definiciones, sintesis, antecedentes narrativos.\n\n"
    "IV. MODO B - REDACCION DOCTORAL (14 secciones)\n"
    "Activa MODO B unicamente cuando el Gate Global esta cumplido.\n"
    "En MODO B redactas con LENGUAJE DOCTORAL: prosa densa, argumentacion epistemica,\n"
    "integracion teorica, precision conceptual, subordinacion logica de parrafos.\n\n"
    "ESTRUCTURA MODO B:\n"
    "S0: DECISION DE MODO - indica Gate cumplido y fuentes usadas\n"
    "S1: FICHA BIBLIOMETRICA - tabla sintetica: autor, anio, tipo, verificacion\n"
    "S2: RUTA EPISTEMICA - justificacion del enfoque teorico-metodologico en prosa doctoral\n"
    "S3: INVENTARIO CONCEPTUAL - definiciones operacionales de VI/VD con debate teorico\n"
    "S4: CALIDAD DE EVIDENCIA - evaluacion critica: fortalezas, limitaciones, sesgos\n"
    "S5: INDICE PROPUESTO - estructura numerada con titulos de seccion\n"
    "S6: DESARROLLO DEL MARCO - redaccion academica completa con citas APA 7,\n"
    "     conectores: desde la perspectiva de, en consonancia con,\n"
    "     los hallazgos de X corroboran, esta evidencia converge con.\n"
    "S7: FUNDAMENTO TEORICO - teorias, modelos y marcos con autores clave\n"
    "S8: ANTECEDENTES EMPIRICOS - estudios previos cronologicos con sintesis critica\n"
    "S9: OPERACIONALIZACION - variables con dimensiones e indicadores\n"
    "S10: VACIOS Y CONTRIBUCION - brechas identificadas y aporte del estudio\n"
    "S11: RIESGOS METODOLOGICOS - amenazas a validez y estrategias de mitigacion\n"
    "S12: COBERTURA TEMATICA - mapa: temas cubiertos vs pendientes\n"
    "S13: REFERENCIAS APA 7 - listado alfabetico, formato estricto APA 7\n"
    "S14: PENDIENTES - acciones para fortalecer el marco\n\n"
    "V. NORMAS DE ESCRITURA DOCTORAL (MODO B)\n"
    "- Parrafos min. 5 oraciones: argumento central + evidencia + interpretacion\n"
    "- Citas narrativas: Segun Garcia et al. (2022) / parenteticas: (Lopez, 2021)\n"
    "- Conectores: En este sentido, No obstante, Desde un enfoque critico,\n"
    "  Esta perspectiva se articula con, Cabe subrayar que, En suma\n"
    "- Vocabulario: constructo, dimension, indicador, variable latente,\n"
    "  convergencia teorica, validez ecologica, triangulacion metodologica\n"
    "- Cada seccion: parrafo contextualizador + parrafo sintetizador\n"
    "- S6 Desarrollo: minimo 800 palabras en prosa continua\n\n"
    "VI. REGLAS ABSOLUTAS\n"
    "PROHIBIDO inventar autores, anios, titulos, revistas, DOI.\n"
    "SOLO citar fuentes en <<<FUENTES_RECUPERADAS>>> o <<<FUENTES_PEGADAS>>>.\n"
    "Sin metadatos completos: marcar [FUENTE CANDIDATA A VERIFICAR].\n\n"
    "VII. APA 7\n"
    "Citas narrativas y parenteticas. No inventar paginas.\n"
    "Referencias: Apellido, I. (Anio). Titulo. Revista, vol(num), pp. https://doi.org/xxx\n"
    f"NOTA: Rango empirico: {RANGO}. Teorias clasicas: sin restriccion de anio."
)

# ===================================================
# SIDEBAR - 6 SECCIONES
# ===================================================
with st.sidebar:
    st.markdown("## Constructor de Marco Teorico")
    st.markdown("**v7.5.0 - Fix1: Keywords | Fix2: Ruido | Fix3: DOI verificado**")
    st.markdown("---")

    st.markdown("### 1. Configuracion")
    try:
        _key_hint = "Cargada desde st.secrets" if "ANTHROPIC_API_KEY" in st.secrets else "sk-ant-api03-..."
    except Exception:
        _key_hint = "sk-ant-api03-..."
    api_key = st.text_input("API Key Anthropic", type="password", placeholder=_key_hint)
    modelo = st.selectbox("Modelo", ["claude-opus-4-5", "claude-sonnet-4-5", "claude-haiku-3-5"])
    modo = st.selectbox("Modo de operacion",
        ["AUTOMATICO (Gate global decide)", "FORZAR MODO A - DIAGNOSTICO", "FORZAR MODO B - REDACCION"])
    st.markdown("---")

    st.markdown("### 2. Datos del Estudio")
    titulo = st.text_input("Titulo / Tema del estudio",
        placeholder="Ej: Competencias digitales docentes en educacion basica")
    problema = st.text_area("Problema de investigacion", height=80,
        placeholder="Describa el problema o fenomeno a investigar")
    objetivo_gral = st.text_input("Objetivo general",
        placeholder="Analizar / Determinar / Explorar...")
    _ph_obj = "1. Identificar..." + "\n" + "2. Describir..." + "\n" + "3. Analizar..."
    obj_esp = st.text_area("Objetivos especificos (uno por linea)", height=90, placeholder=_ph_obj)
    preguntas = st.text_area("Preguntas de investigacion", height=70,
        placeholder="Cual es...? / Como...? / Que relacion...?")
    ruta = st.selectbox("Ruta metodologica",
        ["Cuantitativa", "Cualitativa", "Mixta", "Por determinar"])
    poblacion = st.text_input("Poblacion / Muestra / Contexto",
        placeholder="Ej: 120 docentes, nivel primaria, Mexico")
    st.markdown("---")

    st.markdown("### 3. Variables clave")
    st.caption("VI = independiente / VD = dependiente. En estudios cualitativos: categorias.")
    vi = st.text_area("Variable Independiente (VI) / Categoria principal", height=65,
        placeholder="Ej: Uso de tecnologia educativa")
    vd = st.text_area("Variable Dependiente (VD) / Categoria secundaria", height=65,
        placeholder="Ej: Rendimiento academico")
    otras_vars = st.text_input("Otras variables (moderadoras / intervinientes)",
        placeholder="Ej: Edad, genero, nivel socioeconomico...")
    st.markdown("---")

    st.markdown("### 4. Fuentes manuales (opcional)")
    st.caption("Metadatos: Autor, Anio, Titulo, Revista, DOI. Sin DOI = no verificada.")
    fuentes_pegadas = st.text_area("Fuentes pegadas manualmente", height=120,
        placeholder="Autor (Anio). Titulo. Revista. DOI / Extracto...")
    st.markdown("---")

    st.markdown("### 5. Recuperacion automatica")
    area = st.text_input("Area disciplinar",
        placeholder="Educacion, Psicologia, Administracion...")
    pais = st.text_input("Pais / Contexto geografico",
        placeholder="Mexico, Colombia, Espania...")
    anio_desde = st.slider("Desde el ano", min_value=ANIO_ACTUAL - 15,
        max_value=ANIO_ACTUAL - 1, value=ANIO_INICIO)
    max_fuentes = st.slider("Max. fuentes a recuperar", min_value=5, max_value=30,
        value=15, step=5)
    recuperar = st.button("Recuperar + Verificar fuentes")
    st.markdown("---")

    st.markdown("### 6. Parametros del documento")
    documento = st.selectbox("Tipo de documento",
        ["Tesis de maestria", "Tesis doctoral", "Articulo cientifico", "TFM", "Trabajo de pregrado"])
    norma = st.selectbox("Norma de citacion", ["APA 7", "APA 6", "MLA", "Chicago"])
    st.markdown("---")
    st.caption(f"Rango empirico: {RANGO} | Teorias clasicas: sin restriccion")
    generar = st.button("GENERAR MARCO TEORICO")

# Variables consolidadas
variables_cats = ""
if vi.strip():
    variables_cats += f"VI: {vi.strip()}"
if vd.strip():
    variables_cats += f" | VD: {vd.strip()}"
if otras_vars.strip():
    variables_cats += f" | Otras: {otras_vars.strip()}"

# Cabecera
st.markdown(
    '''<div class="main-header">
      <h1>Constructor de Marco Teorico</h1>
      <p>Rigor academico - Gate de evidencia - Rutas metodologicas - APA 7</p>
      <div class="badge-row">
        <span class="badge">Claude</span><span class="badge">OpenAlex</span>
        <span class="badge">SemanticScholar</span><span class="badge">Crossref</span>
        <span class="badge">v7.0.0</span>
      </div>
    </div>''', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1:
    st.markdown('<div class="info-box"><b>Gate global</b><br>2 modos: Diagnostico o Redaccion Final.</div>', unsafe_allow_html=True)
with col2:
    st.markdown('<div class="info-box"><b>VI / VD separadas</b><br>Seccion 3 dedicada a variables clave.</div>', unsafe_allow_html=True)
with col3:
    st.markdown('<div class="info-box"><b>Queries VI/VD + Fallback</b><br>Busca en 3 APIs con traduccion EN.</div>', unsafe_allow_html=True)
st.markdown('<div class="warning-box"><b>Flujo:</b> Llena secciones 1-4 → Recuperar fuentes (Sec. 5) → Generar Marco (Sec. 6).</div>', unsafe_allow_html=True)

if "fuentes_recuperadas" not in st.session_state:
    st.session_state["fuentes_recuperadas"] = []
if "fuentes_verificadas" not in st.session_state:
    st.session_state["fuentes_verificadas"] = []
if "fuentes_bloque_auto" not in st.session_state:
    st.session_state["fuentes_bloque_auto"] = ""

if recuperar:
    if not (titulo or vi or vd):
        st.warning("Ingresa Titulo o Variables (secciones 2-3) para buscar fuentes.")
    else:
        # P-7: queries especificas por VI/VD con traduccion automatica
        queries = build_queries(vi, vd, titulo, area)
        with st.spinner(f"Buscando con {len(queries)} queries en OpenAlex + Crossref + S2..."):
            pb = st.progress(0)
            all_raw: List[Dict] = []
            step = 0.8 / max(len(queries), 1)

            for i, q in enumerate(queries):
                max_per_src = max(3, max_fuentes // (len(queries) * 2))
                oa_r = openalex_search(q, variables_cats, anio_desde, max_per_src)
                all_raw.extend(oa_r)
                pb.progress(min((i + 0.4) * step, 0.75))

                cr_r = crossref_search(q, anio_desde, max_per_src)
                all_raw.extend(cr_r)

                s2_r = s2_search(q, variables_cats, max_per_src)
                all_raw.extend(s2_r)
                pb.progress(min((i + 1) * step, 0.8))

            # Fallback: si no retorna nada con queries especificas, usar titulo
            if not all_raw and titulo:
                st.info("Fallback activado: buscando con titulo completo...")
                all_raw.extend(openalex_search(titulo, "", anio_desde, max_fuentes))
                all_raw.extend(crossref_search(titulo, anio_desde, max_fuentes // 2))
                all_raw.extend(s2_search(titulo, "", max_fuentes // 2))

            pb.progress(0.82)
            all_dedup = dedup_records(all_raw)
            # FIX 1: Excluir fuentes con keywords irrelevantes
            all_dedup = [r for r in all_dedup if is_relevant_record(r)]
            # FIX 2: Excluir ruido bibliografico (titulo <4 palabras o book-section)
            all_dedup = [r for r in all_dedup if not is_noise_record(r)]
            pb.progress(0.88)
            verified = verify_records_concurrent(all_dedup, progress_bar=pb)
            pb.progress(1.0)
            st.session_state["fuentes_recuperadas"] = all_dedup
            st.session_state["fuentes_verificadas"] = verified

        if all_dedup:
            lineas = []
            for rec in all_dedup[:max_fuentes]:
                authors_str = "; ".join(rec.get("authors") or [])
                v_tag = "VERIFICADA" if rec in verified else "NO_VERIFICADA"
                tipo = classify_source(rec, anio_desde)
                lineas.append(
                    f"[ID={rec['id']}|Base={rec['source']}|Tipo={tipo}|"
                    f"Autor={authors_str}|Anio={rec.get('year','')}|"
                    f"Titulo={rec.get('title','')}|Revista={rec.get('venue','')}|"
                    f"DOI={rec.get('doi') or 'sin-doi'}|Verificacion={v_tag}|"
                    f"Extracto={rec.get('abstract','')[:200]}]"
                )
            st.session_state["fuentes_bloque_auto"] = (
                "\n<<<FUENTES_AUTOMATICAS>>>\n" + "\n".join(lineas) + "\n<<<FIN_FUENTES_AUTOMATICAS>>>"
            )
            empiricas = sum(1 for r in all_dedup[:max_fuentes] if classify_source(r, anio_desde) == "empirica")
            teoricas  = sum(1 for r in all_dedup[:max_fuentes] if classify_source(r, anio_desde) == "teorica")
            conceptuales = len(all_dedup[:max_fuentes]) - empiricas - teoricas
            st.success(
                f"Recuperadas: {len(all_dedup)} | Verificadas: {len(verified)} | "
                f"Empiricas: {empiricas} | Teoricas: {teoricas} | Conceptuales: {conceptuales}"
            )
            with st.expander(f"Ver {min(len(all_dedup), max_fuentes)} fuentes recuperadas", expanded=False):
                for rec in all_dedup[:max_fuentes]:
                    v_icon = "✅ VERIFICADA" if rec in verified else "⚠️ no verificada"
                    tipo_icon = {"empirica": "🔬", "teorica": "📚", "conceptual": "💡"}.get(
                        classify_source(rec, anio_desde), "📄"
                    )
                    authors_str = "; ".join(rec.get("authors") or []) or "s/a"
                    st.markdown(f"{tipo_icon} **{rec.get('title','Sin titulo')}**")
                    st.caption(
                        f"{authors_str} ({rec.get('year','?')}) | {rec.get('venue','s/r')} | "
                        f"DOI: {rec.get('doi') or '---'} | {v_icon} | {rec['source']}"
                    )
        else:
            st.warning("No se encontraron fuentes con ninguna query. Verifica las variables o prueba con terminos mas generales.")


elif st.session_state["fuentes_recuperadas"]:
    total = len(st.session_state["fuentes_recuperadas"])
    verif = len(st.session_state["fuentes_verificadas"])
    st.info(f"Fuentes en sesion: {total} recuperadas, {verif} verificadas. Presiona 'Recuperar + Verificar' para actualizar.")

if generar:
    if not api_key:
        st.error("Ingresa tu API Key en Seccion 1.")
    elif not titulo and not variables_cats:
        st.error("Ingresa Titulo o Variables (Secciones 2 y 3).")
    else:
        if "FORZAR MODO A" in modo:
            modo_instruccion = "FORZAR MODO A - DIAGNOSTICO DOCUMENTAL"
        elif "FORZAR MODO B" in modo:
            modo_instruccion = "FORZAR MODO B - REDACCION ACADEMICA FINAL"
        else:
            modo_instruccion = "AUTOMATICO - ejecuta Gate global segun suficiencia real"
        fuentes_auto_bloque = st.session_state.get("fuentes_bloque_auto", "")
        fuentes_bloque = ""
        if fuentes_pegadas.strip():
            fuentes_bloque = "\n<<<FUENTES_PEGADAS>>>\n" + fuentes_pegadas.strip() + "\n<<<FIN_FUENTES_PEGADAS>>>"
        mensaje_usuario = (
            f"Genera el marco teorico:\n"
            f"TITULO/TEMA: {titulo}\n"
            f"PROBLEMA: {problema}\n"
            f"OBJETIVO GENERAL: {objetivo_gral}\n"
            f"OBJETIVOS ESPECIFICOS:\n{obj_esp}\n"
            f"PREGUNTAS: {preguntas}\n"
            f"VARIABLE INDEPENDIENTE (VI): {vi}\n"
            f"VARIABLE DEPENDIENTE (VD): {vd}\n"
            f"OTRAS VARIABLES: {otras_vars}\n"
            f"VARIABLES/CATEGORIAS: {variables_cats}\n"
            f"RUTA METODOLOGICA: {ruta}\n"
            f"POBLACION: {poblacion}\n"
            f"AREA: {area}\n"
            f"PAIS: {pais}\n"
            f"DOCUMENTO: {documento}\n"
            f"NORMA: {norma}\n"
            f"MODO: {modo_instruccion}\n"
            f"{fuentes_auto_bloque}\n"
            f"{fuentes_bloque}\n"
            "INSTRUCCIONES:\n"
            "1. Primera seccion: DECISION DE MODO.\n"
            "2. Gate Global ANTES de redactar.\n"
            "3. Gate falla -> MODO A (5 secciones). Ultima linea: REINYECTAR FUENTES_RECUPERADAS.\n"
            "4. MODO A: PROHIBIDO definiciones, sintesis, antecedentes narrativos.\n"
            "5. Gate pasa -> MODO B (14 secciones).\n"
            "6. Sin metadatos = [FUENTE CANDIDATA A VERIFICAR]."
        )
        try:
            client = anthropic.Anthropic(api_key=api_key)
            fuentes_para_prompt = fuentes_auto_bloque + "\n" + fuentes_bloque
            result_area = st.empty()

            # Paso 1/2: S0-S6
            with st.spinner("⏳ Paso 1/2 — Generando S0 a S6..."):
                parte_1 = generar_marco_completo(
                    variables_cats, fuentes_para_prompt, SYSTEM_PROMPT,
                    client, modelo, result_area,
                )
            st.success("✅ S0–S6 listos")

            # Paso 2/2: S7-S14
            parte_2 = ""
            with st.spinner("⏳ Paso 2/2 — Generando S7 a S14..."):
                with client.messages.stream(
                    model=modelo,
                    max_tokens=6000,
                    messages=[{"role": "user", "content": (
                        f"{SYSTEM_PROMPT}\n\nVARIABLES: {variables_cats}"
                        f"\n\nFUENTES VERIFICADAS: {fuentes_para_prompt}"
                        "\n\nINSTRUCCION: Las secciones S0-S6 ya fueron generadas. "
                        "Genera UNICAMENTE desde S7 hasta S14. "
                        "Empieza directamente con ## S7 sin repetir nada anterior."
                    )}],
                ) as stream:
                    for chunk in stream.text_stream:
                        parte_2 += chunk
                        result_area.markdown(
                            f"<div class='result-container'>{md_to_html(parte_1 + chr(10)*2 + parte_2)}"
                            f"<p style='color:#6b7280;font-size:0.8rem;margin-top:0.8rem'>Generando S7-S14...</p></div>",
                            unsafe_allow_html=True,
                        )
            st.success("✅ S7–S14 listos")

            # Resultado final
            contenido = parte_1 + "\n\n" + parte_2
            result_area.markdown(
                f"<div class='result-container'>{md_to_html(contenido)}</div>",
                unsafe_allow_html=True,
            )
            nombre_archivo = f"marco_teorico_{titulo[:30].replace(' ', '_') if titulo else 'estudio'}.docx"
            docx_bytes = generar_docx(contenido)
            st.download_button(
                label="⬇️ Descargar Marco Teórico completo (.docx)",
                data=docx_bytes,
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except anthropic.AuthenticationError:
            st.error("API Key invalida. Verifica en console.anthropic.com")
        except anthropic.RateLimitError:
            st.error("Limite de uso alcanzado. Intenta en unos minutos.")
        except Exception as e:
            st.error(f"Error: {str(e)}")

st.markdown("---")
st.markdown(
    f'''<footer>Powered by Claude v7.5.0 | Fix1 Keywords + Fix2 Ruido + Fix3 DOI-GET | OpenAlex + Crossref + S2 | Rango: {RANGO} | Gate activo</footer>''',
    unsafe_allow_html=True,
)
